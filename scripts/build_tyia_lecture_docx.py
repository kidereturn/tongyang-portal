# -*- coding: utf-8 -*-
"""
TYIA 사이트 설명 대본 docx 생성
- 대상: IT 지식이 거의 없는 일반 임직원
- 톤: 친근함 + 구체적 + 비유 사용 (전문용어 최소화, 사용 시 즉시 풀어 설명)
- 권오윤 차장 강의 대본 톤 참고 (1차 강의 PDF + 본문 일부)
"""
import os, sys
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, 'docs', 'TYIA_설명대본_일반임직원.docx')

doc = Document()

# 기본 폰트
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(11)

# 페이지 여백
section = doc.sections[0]
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── 헬퍼 ───
def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor(0x0F, 0x1E, 0x36)
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15); r.font.bold = True
    r.font.color.rgb = RGBColor(0x31, 0x82, 0xF6)
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x2E, 0x4D)
    return p

def add_para(text, *, bold=False, italic=False, color=None, size=11, indent=False, line_space=1.5):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = line_space
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_quote(text):
    """ 강조 인용 (대본 발언 부분) """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"💬  {text}")
    r.font.size = Pt(11); r.font.italic = True
    r.font.color.rgb = RGBColor(0x4E, 0x5A, 0x6F)
    return p

def add_bullet(text, *, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + 0.5 * level)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def add_callout(title, body, *, color=(0x31, 0x82, 0xF6)):
    """ 박스 형태 콜아웃 (1행 1열 표) """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    # 배경색
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF4FE')
    tcPr.append(shd)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.font.size = Pt(11); r1.font.bold = True
    r1.font.color.rgb = RGBColor(*color)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x1A, 0x2E, 0x4D)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()  # spacing

def add_term(term, desc):
    """ 용어 풀이 (term: 설명) """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"• {term}")
    r1.font.size = Pt(11); r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x31, 0x82, 0xF6)
    r2 = p.add_run(f"  —  {desc}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x4E, 0x5A, 0x6F)

def add_separator():
    p = doc.add_paragraph()
    r = p.add_run('—' * 40)
    r.font.color.rgb = RGBColor(0xB0, 0xB8, 0xC1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════
add_h1('TYIA — 내부회계 포털 사용 설명회')
add_para('대본: IT 지식이 거의 없는 일반 임직원 대상', italic=True, color=RGBColor(0x8B, 0x95, 0xA1))
add_para('(주)동양 내부회계관리실 · 2026년 4월', italic=True, color=RGBColor(0x8B, 0x95, 0xA1))
add_para('')
add_para('전체 발표 시간: 약 25분 (시연 포함) + Q&A 10분', size=10, color=RGBColor(0x4E, 0x5A, 0x6F))

add_callout(
    '📍 이 자료를 보시는 분께',
    '이 문서는 발표 대본입니다. 실제로 발표할 때 그대로 읽어도 어색하지 않게 일상 대화체로 작성했습니다. '
    '전문용어가 처음 나올 때마다 풀어서 설명드리고, 가능한 한 쉬운 비유를 사용했습니다. '
    '굵은 글씨는 핵심 메시지, 💬 표시는 직접 발화하시는 부분입니다.'
)

add_separator()

# ════════════════════════════════════════════
# 1. 인사 & 도입
# ════════════════════════════════════════════
add_h1('1. 시작 인사')

add_para('(가벼운 자기소개로 분위기 만들기)')
add_quote(
    '안녕하세요. 내부회계관리실에서 나왔습니다. '
    '오늘 소개해드릴 시스템은 ‘TYIA’ 라고 부르는 내부회계 포털입니다. '
    '사이트 주소는 tyia.vercel.app 한 줄만 외워두시면 됩니다.'
)
add_quote(
    '시작에 앞서 한 가지 약속드릴게요. '
    '오늘 이 발표를 끝까지 들으시면, 여러분 자리로 돌아가셔서 바로 사용할 수 있도록 '
    '쉽게, 그리고 천천히 설명드리겠습니다. 모르는 단어가 나오면 그 자리에서 바로 풀어드릴게요.'
)

add_h2('1-1. 왜 이 시스템이 필요한가요?')
add_para(
    '이전에는 통제활동 점검 자료를 모을 때, 각자 PC 폴더에 정리하거나 메일로 받고, '
    '엑셀에 다시 정리하는 번거로움이 있었습니다. 부서끼리 양식이 다르고, 누가 언제 제출했는지 '
    '추적도 어려웠죠.'
)
add_quote(
    '그래서 “이 모든 걸 한 곳에서, 클릭 한두 번으로 끝낼 수 있으면 좋겠다”는 생각으로 만든 것이 '
    '오늘 보여드릴 TYIA 입니다.'
)

add_separator()

# ════════════════════════════════════════════
# 2. 핵심 한 줄
# ════════════════════════════════════════════
add_h1('2. 한 줄로 정리하면')

add_callout(
    '핵심',
    '“증빙 첨부, 결재상신, 학습 — 모두 한 곳에서.”',
    color=(0x31, 0x82, 0xF6),
)

add_para('이걸 풀어 말씀드리면 세 단계입니다:', bold=True)
add_bullet('① 매일 1번 — 포털에 들어와서 본인이 담당하는 통제활동을 확인합니다.')
add_bullet('② 증빙 파일을 — 마우스로 끌어다 놓기만 하면 자동으로 저장됩니다.')
add_bullet('③ 결재상신을 — 한 번 클릭하면 끝납니다.')

add_quote(
    '복잡한 시스템 아닙니다. 핵심은 세 단계예요. 접속, 업로드, 상신. 이게 다입니다. '
    '실제로 처음 사용하시는 분도 5분이면 충분합니다.'
)

add_separator()

# ════════════════════════════════════════════
# 3. 로그인
# ════════════════════════════════════════════
add_h1('3. 로그인 — 이것만 외우세요')

add_h2('3-1. 사이트 접속')
add_para('인터넷 익스플로러는 안 됩니다. ', bold=True, color=RGBColor(0xDC, 0x26, 0x26))
add_para('크롬 / 엣지 / 사파리 / 네이버 웨일 — 이 중 아무거나 사용하시면 됩니다.')

add_callout('주소', 'https://tyia.vercel.app')

add_h2('3-2. 사번과 비밀번호')
add_para('처음 들어오시면 로그인 화면이 나옵니다. 두 칸만 채우시면 됩니다.', bold=True)
add_term('사번', '회사에서 받은 6자리 숫자 (예: 101267)')
add_term('비밀번호', '“ty” + 사번 (예: ty101267) — 영문 ty 는 모두 소문자입니다')

add_quote(
    '예를 들어 사번이 101267 이면, 비밀번호는 ty101267 입니다. '
    '첫 로그인 후에는 본인이 원하는 비밀번호로 바꿀 수 있어요. '
    '꼭 한번 바꿔주시기 바랍니다 — 보안 문제거든요.'
)

add_callout(
    '🔐 보안 안내',
    '회사 PC를 다른 분과 같이 쓰시거나, 회의실 PC 같은 공용 PC 에서 로그인하셨으면 '
    '반드시 사용 후 로그아웃하셔야 합니다. 로그아웃은 우측 상단 본인 이름 클릭 후 보이는 메뉴에 있습니다.',
    color=(0xDC, 0x26, 0x26),
)

add_separator()

# ════════════════════════════════════════════
# 4. 역할 안내
# ════════════════════════════════════════════
add_h1('4. 본인 역할 확인')

add_para(
    '로그인 하시면 화면 우측 상단에 본인 이름과 함께 영어 한 단어가 표시됩니다. '
    '이 단어가 여러분의 역할입니다. 세 가지 중 하나입니다:', bold=True
)

# 표
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['역할', '한국어', '인원', '하시는 일']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs: r.font.bold = True; r.font.size = Pt(11)
data = [
    ('OWNER',      '담당자',  '442명', '본인 통제활동의 증빙을 업로드 + 결재상신'),
    ('CONTROLLER', '승인자',  '33명',  '담당자가 올린 증빙을 검토 + 승인 / 반려'),
    ('ADMIN',      '관리자',  '3명',   '내부회계팀 — 전사 관리 + 검토 메모'),
]
for i, (a, b, c, d) in enumerate(data, 1):
    cells = table.rows[i].cells
    cells[0].text = a; cells[1].text = b; cells[2].text = c; cells[3].text = d
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(10)

add_para('')
add_quote(
    '본인이 어떤 역할인지 모르시겠으면 우측 상단을 보세요. '
    'OWNER, CONTROLLER, ADMIN 중 하나가 표시됩니다. '
    '대부분의 직원분은 OWNER 입니다. 그게 가장 많거든요. 442명이세요.'
)

add_separator()

# ════════════════════════════════════════════
# 5. 담당자 사용법 (가장 중요)
# ════════════════════════════════════════════
add_h1('5. 담당자(OWNER)가 하실 일 — 가장 중요합니다')

add_para(
    '대부분이 담당자이시니까 여기를 가장 길게 설명드리겠습니다. '
    '단계별로 하나씩 보여드릴게요.', bold=True
)

add_h2('Step 1. 상단 메뉴에서 “증빙관리” 클릭')
add_quote(
    '로그인 하면 홈 화면이 보이고요. 상단에 메뉴 7~8개가 가로로 늘어서 있습니다. '
    '그 중에 “증빙관리” 라는 메뉴를 클릭하시면 됩니다.'
)

add_h2('Step 2. 본인 담당 통제활동 목록 확인')
add_para('증빙관리 메뉴를 누르면 본인이 담당하는 통제활동 목록이 표로 보입니다.', bold=True)
add_para('이 표에는 다음 정보가 있습니다:')
add_bullet('통제번호 — TR.05.W2.C2 같은 코드 (외울 필요 없습니다)')
add_bullet('통제활동명 — 무엇을 점검하는지 (예: 신규 자금 조달 검토 및 승인)')
add_bullet('담당부서 / 담당자 / 승인자 — 자동으로 매핑되어 있습니다')
add_bullet('건수 — 업로드해야 할 모집단 / 이미 업로드한 행 (예: 4/10)')
add_bullet('KPI — 통제활동의 점수와 등급 (관리자가 평가)')
add_bullet('상신 상태 — 미완료 / 상신완료 / 승인완료 / 수정제출')

add_callout(
    '💡 모집단이 뭐예요?',
    '“모집단”은 점검 대상 거래·문서 목록입니다. 예를 들어 “11월 자금 결재 100건” 이라면 '
    '모집단은 100건이고, 그 중 정해진 샘플(예: 10건)에 대해 증빙을 첨부합니다. '
    '내부회계팀에서 미리 작성해두기 때문에 여러분이 만드실 필요는 없습니다.'
)

add_h2('Step 3. 행마다 있는 “증빙확인” 버튼 클릭')
add_quote(
    '각 행 우측 끝에 “증빙확인”이라는 파란색 버튼이 있어요. 그걸 누르시면 '
    '팝업창이 하나 떠요. 그 안에서 모든 작업을 하시는 겁니다.'
)

add_h2('Step 4. 팝업창 — 모집단 별로 파일 업로드')
add_para(
    '팝업창이 열리면 가운데에 표가 있고, 표의 각 행이 “모집단 행” 입니다. '
    '한 행에 한 개 또는 여러 개의 증빙 파일을 첨부할 수 있어요.', bold=True
)

add_para('파일 첨부 방법은 두 가지입니다:')
add_bullet('① 파일을 마우스로 끌어다가 우측의 박스 안에 놓기 (드래그앤드롭)')
add_bullet('② 박스를 클릭해서 윈도우 파일 탐색기에서 파일 선택')

add_callout(
    '✨ 여기가 신기한 부분',
    '파일을 하나 추가할 때마다 자동으로 “중간저장” 됩니다. '
    '“저장 안 해서 잃어버릴까봐” 걱정하실 필요 없습니다. '
    '브라우저를 갑자기 닫아도, 컴퓨터가 재부팅 되어도 — 파일은 안 사라집니다.'
)

add_h2('Step 5. 잘못 올렸으면 — 교체하거나 삭제')
add_para('업로드된 파일 옆에는 세 가지 버튼이 있습니다:')
add_term('다운로드', '본인이 올린 파일을 다시 받아 확인')
add_term('교체', '같은 자리에 다른 파일로 바꾸기')
add_term('삭제', '파일을 빼기 (그 자리는 다시 비어집니다)')

add_h2('Step 6. 모든 모집단 완료 후 — 결재상신')
add_para(
    '모든 모집단 행에 파일이 들어가면 (예: 10/10), 우측 하단의 ', bold=True
)
add_para('“결재상신” 버튼이 활성화', bold=True, color=RGBColor(0x31, 0x82, 0xF6))
add_para('됩니다. 그걸 누르시면 끝입니다. 승인자에게 자동으로 알림이 갑니다.')

add_quote(
    '한 가지만 기억하세요. 결재상신은 모든 모집단이 다 채워져야 활성화됩니다. '
    '하나라도 비어있으면 버튼이 회색으로 눌리지 않습니다. 그건 “아직 다 안 끝났어요” 라는 뜻이에요.'
)

add_callout(
    '⏰ 마감일은?',
    '매 분기 마감일은 별도로 공지됩니다. 보통 분기 종료 후 2주 안입니다. '
    '홈 화면 하단에 “월말 마감 D-day” 가 자동 표시되니까 참고하세요.',
    color=(0xF5, 0x9E, 0x0B),
)

add_separator()

# ════════════════════════════════════════════
# 6. 자주 묻는 질문 (담당자)
# ════════════════════════════════════════════
add_h1('6. 자주 묻는 질문 — 담당자')

add_h3('Q1. 잘못 올렸어요. 어떻게 바꾸나요?')
add_para('A. 파일 옆의 “교체” 버튼 → 새 파일 선택 → 자동 중간저장 → 끝.')

add_h3('Q2. 삭제하고 다시 올리고 싶어요.')
add_para('A. “삭제” 버튼 클릭 → 즉시 시스템에서 빠집니다. 그 후 다시 올리시면 됩니다.')

add_h3('Q3. 상신 후에도 수정할 수 있나요?')
add_para('A. 두 가지 경우가 있어요:')
add_bullet('승인 전: 승인자가 “반려”하시면 다시 수정 가능합니다.')
add_bullet('승인 후: 관리자(내부회계팀)에게 “수정제출” 을 요청하셔야 해요. 그러면 다시 미완료 상태로 돌아옵니다.')

add_h3('Q4. 다른 사람의 증빙도 볼 수 있나요?')
add_para('A. ❌ 절대 불가능합니다. ', bold=True, color=RGBColor(0xDC, 0x26, 0x26))
add_para(
    '시스템적으로 본인 담당 활동만 보이도록 설계되어 있어요. '
    '“RLS” 라는 데이터베이스 보안 정책으로 막혀 있습니다.'
)

add_h3('Q5. PC 가 멈췄어요. 작업 데이터 다 날아간 건가요?')
add_para(
    'A. 안 날아갑니다. 모든 파일은 업로드 즉시 회사 클라우드 서버에 저장됩니다. '
    '브라우저를 종료해도, 컴퓨터를 끄셔도, 며칠 후에 다시 들어오시면 그대로 있습니다.'
)

add_h3('Q6. 사용 중에 문제가 생기면 누구한테 물어보나요?')
add_para('세 가지 방법이 있어요:')
add_bullet('① 사이트 안의 “AI 챗봇” — 24시간 즉답 (오른쪽 상단 메뉴)')
add_bullet('② “Tell me” 게시판 — 익명으로 의견 남기기')
add_bullet('③ 내부회계팀 — 02-6150-7000 (평일 09:00~18:00)')

add_separator()

# ════════════════════════════════════════════
# 7. 승인자 사용법
# ════════════════════════════════════════════
add_h1('7. 승인자(CONTROLLER)가 하실 일')

add_para('승인자는 33명입니다. 보통 팀장님급이세요.', italic=True)

add_h2('Step 1. 본인 담당 통제활동 확인')
add_para('승인자가 “증빙관리” 들어가시면 — 담당자분과 다른 화면이 보입니다.', bold=True)
add_bullet('본인이 검토하실 통제활동만 표시됩니다 (= 본인이 “승인자”로 지정된 활동들)')
add_bullet('각 행에 “결재” 컬럼이 있고, 거기에 ‘승인’ 또는 ‘반려’ 버튼이 보이면 처리할 항목입니다')

add_h2('Step 2. 모달 안에서 파일 검토')
add_para('“증빙확인”을 누르면 담당자가 올린 모든 파일을 보실 수 있어요.')
add_bullet('파일 별로 “다운로드” 버튼이 있어서 본인 PC에 받아 확인 가능')
add_bullet('“전체 ZIP 다운로드” 버튼으로 한 번에 다 받기 가능')

add_h2('Step 3. 승인 또는 반려')

add_callout(
    '✅ 승인',
    '문제가 없으면 “승인” 버튼 클릭. 즉시 “승인완료” 상태로 전환되고, 담당자에게 알림 갑니다.',
    color=(0x16, 0xA3, 0x4A),
)
add_callout(
    '↩️ 반려',
    '수정이 필요하면 “반려” 클릭 → 사유 입력창이 떠요 → 사유 적고 확인. '
    '담당자가 즉시 알림 받고, 다시 수정해서 재상신할 수 있습니다.',
    color=(0xDC, 0x26, 0x26),
)

add_quote(
    '반려하실 때 한 가지 부탁드릴 게 있어요. '
    '사유를 꼭 구체적으로 적어주세요. “파일 다시” 이런 식으로 적으시면 담당자가 무엇을 수정해야 할지 모릅니다. '
    '“1번 모집단의 결재서류가 누락됐습니다” 처럼 적어주시면 한 번에 처리됩니다.'
)

add_separator()

# ════════════════════════════════════════════
# 8. 부가 기능
# ════════════════════════════════════════════
add_h1('8. 그 외 부가 기능 — 시간 나실 때 한번 써보세요')

add_para('내부회계 외에도 학습과 소통을 위한 기능이 있습니다.', bold=True)

add_h3('🎯 빙고게임')
add_para('매일 25문제. 1줄을 완성하거나, 오답 1번이면 도전 1회를 사용하고, 일 최대 3회까지 가능합니다. 정답을 맞히면 포인트가 적립됩니다.')

add_h3('🎓 강좌')
add_para('내부회계 교육 영상이 올라와 있어요. 영상 시청 후 간단한 퀴즈를 풀면 수료 + 포인트. 강사에게 익명으로 질문도 할 수 있습니다.')

add_h3('📰 회사소식')
add_para('동양 공시(DART) + 동양 주가 + 레미콘/건설/플랜트 뉴스를 자동으로 모아서 보여드립니다. 매일 한 번 갱신됩니다.')

add_h3('🤖 AI 챗봇')
add_para('“통제활동 TR.05.W2.C2 가 뭐야?” 처럼 자연어로 물어보시면 답변해줍니다. 24시간 사용 가능.')

add_h3('💬 Tell me')
add_para('회사에 대한 의견·불편·건의를 익명으로 남기실 수 있어요. 직책 상관없이 자유롭게 적으세요. (관리자만 신원 확인 가능 — 운영 목적)')

add_h3('🎨 웹툰')
add_para('내부회계 관련 짧은 만화. 댓글도 익명으로 가능합니다.')

add_h3('🏅 포인트 / 랭킹')
add_para('활동별로 포인트가 쌓입니다. 월말 TOP 10 은 별도 시상이 있을 수 있어요 (운영진 결정).')

add_separator()

# ════════════════════════════════════════════
# 9. 도움이 필요할 때
# ════════════════════════════════════════════
add_h1('9. 막히면 — 이 4가지 중 하나')

add_callout('🔔 알림함',     '첫 화면 우측 상단 종(벨) 아이콘. 결재 상태 변경, 메모, 공지가 다 모입니다.')
add_callout('🤖 AI 챗봇',    '/chatbot — 24시간 즉답. 가벼운 질문은 여기로.', color=(0x16, 0xA3, 0x4A))
add_callout('💬 Tell me',    '/tellme — 익명 문의. 답변은 운영진이 정리해서 답변합니다.', color=(0xF5, 0x9E, 0x0B))
add_callout('☎️  내부회계팀','02 - 6150 - 7000 (평일 09:00~18:00). 긴급한 건 전화가 가장 빠릅니다.', color=(0xDC, 0x26, 0x26))

add_separator()

# ════════════════════════════════════════════
# 10. 마무리
# ════════════════════════════════════════════
add_h1('10. 마무리')

add_quote(
    '오늘 이 정도 보셨으면 사용에 전혀 지장 없으실 거예요. '
    '한 번 더 정리해드릴게요.'
)

add_para('이것만 외우세요:', bold=True, color=RGBColor(0x31, 0x82, 0xF6))
add_bullet('① 주소: tyia.vercel.app')
add_bullet('② 로그인: 사번 + “ty + 사번”')
add_bullet('③ 핵심: 증빙관리 → 증빙확인 → 드래그앤드롭 → 결재상신')
add_bullet('④ 막히면: AI 챗봇 / Tell me / 02-6150-7000')

add_quote(
    '이상으로 사용법 안내를 마치겠습니다. 지금부터 실제 시연 보여드린 다음 질문 받겠습니다.'
)

add_separator()

# ════════════════════════════════════════════
# 부록: 시연 시나리오
# ════════════════════════════════════════════
add_h1('부록 A. 시연 시나리오 (Live Demo · 5분)')

demo_table = doc.add_table(rows=10, cols=3)
demo_table.style = 'Light Grid Accent 1'
demo_hdr = demo_table.rows[0].cells
for i, h in enumerate(['#', '단계', '예상 시간']):
    demo_hdr[i].text = h
    for p in demo_hdr[i].paragraphs:
        for r in p.runs: r.font.bold = True; r.font.size = Pt(11)
demo_data = [
    ('01', '로그인 (101267 최해성으로)',                '30초'),
    ('02', '홈 화면 둘러보기 (KPI / 검토 상태)',        '30초'),
    ('03', '증빙관리 → 활동 1건 클릭 → 모달 열기',     '30초'),
    ('04', 'PDF 1개 드래그앤드롭 업로드 → 자동저장',   '30초'),
    ('05', '추가 파일 업로드 → 결재상신',              '30초'),
    ('06', '관리자 계정으로 전환 → 검토 → 메모',       '60초'),
    ('07', '빙고게임 한 번 풀어보기',                  '30초'),
    ('08', 'AI 챗봇 1회 질문',                         '30초'),
    ('09', '로그아웃',                                 '15초'),
]
for i, row in enumerate(demo_data, 1):
    cells = demo_table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val
        for p in cells[j].paragraphs:
            for r in p.runs: r.font.size = Pt(10)

add_para('')
add_separator()

# ════════════════════════════════════════════
# 부록 B: 용어 빠른 사전
# ════════════════════════════════════════════
add_h1('부록 B. 용어 빠른 사전')
add_para('이 자료에 나온 용어를 한 줄로 풀어드립니다:', italic=True)

terms = [
    ('TYIA',          '동양 내부회계 포털 시스템 이름'),
    ('통제활동',       '내부회계관리제도에서 정의한 부서별 점검 항목'),
    ('모집단',         '점검 대상이 되는 거래·문서 목록 (내부회계팀에서 미리 작성)'),
    ('증빙',           '통제 수행을 입증하는 PDF·Excel 파일'),
    ('상신',           '담당자가 승인자에게 결재 요청을 보내는 동작'),
    ('승인 / 반려',    '승인자가 결재 처리하는 두 가지 결과'),
    ('수정제출',       '관리자가 승인된 항목을 다시 수정하라고 되돌리는 처리'),
    ('OWNER',         '담당자 (442명)'),
    ('CONTROLLER',    '승인자 (33명)'),
    ('ADMIN',         '관리자 = 내부회계팀 (3명)'),
    ('드래그앤드롭',   '마우스로 파일을 끌어다 박스에 놓는 동작'),
    ('자동 중간저장',  '파일을 추가할 때마다 자동으로 저장하는 기능 — 사용자가 수동으로 저장 안 해도 됨'),
    ('알림함',         '본인에게 온 시스템 알림이 모이는 곳 (우측 상단 벨 아이콘)'),
    ('KPI',           '통제활동의 점수 (90+ A · 85+ B · 80+ C · 75+ D · 70+ E · 그 미만 F)'),
    ('포인트',         '활동마다 적립되는 점수 (로그인 / 빙고 / 강좌 / 댓글 / Tell me)'),
    ('익명 게시',      '신원을 숨기고 글을 쓰는 기능 (관리자만 운영 목적으로 신원 확인 가능)'),
    ('RLS',           '데이터베이스 보안 정책. 시스템적으로 본인 데이터만 보이게 차단함'),
]
for term, desc in terms:
    add_term(term, desc)

doc.save(OUT)
print(f"[OK] saved: {OUT}".encode('ascii', 'replace').decode('ascii'))
print(f"     paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
