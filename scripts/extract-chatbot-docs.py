# -*- coding: utf-8 -*-
"""
Extract text from 19 chatbot source files → JSON for Supabase insert.

Output: scripts/chatbot-docs-extracted.json
   list of { filename, category, title, content }
"""
import os, json, re, sys
from pathlib import Path

SOURCE_DIR = Path(r"C:\Users\tyinc\Downloads\챗봇_소스")
OUT = Path(__file__).parent / "chatbot-docs-extracted.json"

try:
    import docx
except ImportError:
    print("Install: pip install python-docx", file=sys.stderr); sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("Install: pip install openpyxl", file=sys.stderr); sys.exit(1)


def extract_docx(path: Path) -> str:
    """Extract text from docx including tables."""
    d = docx.Document(path)
    chunks = []
    # Paragraphs
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    # Tables
    for table in d.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Avoid dup text from merged cells — but join for now
                cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                chunks.append(' | '.join(row_cells))
    text = '\n'.join(chunks)
    # Normalize whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def extract_xlsx(path: Path) -> str:
    """Extract cell text from xlsx sheets."""
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    chunks = []
    for sheet in wb.worksheets:
        chunks.append(f"[시트] {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            vals = [str(v) for v in row if v not in (None, '')]
            if vals:
                chunks.append(' | '.join(vals))
    text = '\n'.join(chunks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# Classify category by filename keyword
def infer_category(name: str) -> str:
    n = name.replace(' ', '')
    if 'KPI' in n: return 'KPI'
    if '평가및보고가이드라인' in n: return '제도 개요'
    if '감사위원회규정' in n: return '내부통제'
    if '내부감사규정' in n: return '내부통제'
    if '내부신고제도' in n: return '내부통제'
    if '내부회계관리업무지침' in n: return '프레임워크'
    if '내부회계관리규정' in n: return '프레임워크'
    if '기안협조문서' in n: return '동양 특화'
    if '상법시행령' in n: return '법규'
    if '상법' in n and '법률' in n: return '법규'
    if '연간업무일정' in n: return '동양 특화'
    if '외부감사' in n and '회계등에관한규정' in n: return '법규'
    if '외부감사인선임규정' in n: return '감사 개념'
    if '자본시장과금융투자업' in n and '시행규칙' in n: return '법규'
    if '자본시장과금융투자업' in n and '시행령' in n: return '법규'
    if '자본시장과금융투자업' in n and '법률' in n: return '법규'
    if '외부감사에관한법률시행규칙' in n: return '법규'
    if '외부감사에관한법률시행령' in n: return '법규'
    if '외부감사에관한법률' in n: return '법규'
    return '문서'


# Build short title from filename (strip extension, version markers)
def make_title(name: str) -> str:
    base = Path(name).stem
    # Remove trailing (2025.03.10. 개정) tails → keep but shorten
    return base


def main():
    if not SOURCE_DIR.exists():
        print(f"Source dir not found: {SOURCE_DIR}", file=sys.stderr); sys.exit(1)

    files = sorted(SOURCE_DIR.iterdir())
    results = []
    errors = []

    for f in files:
        if f.name.startswith('~$') or f.name.startswith('.'):
            continue
        try:
            suffix = f.suffix.lower()
            if suffix == '.docx':
                text = extract_docx(f)
            elif suffix == '.xlsx':
                text = extract_xlsx(f)
            else:
                print(f"  skip (unsupported): {f.name}")
                continue
            if not text:
                raise RuntimeError('empty extracted text')
            title = make_title(f.name)
            category = infer_category(f.name)
            results.append({
                'filename': f.name,
                'title': title,
                'category': category,
                'content': text,
                'char_count': len(text),
            })
            print(f"  ok  {f.name:60s}  {category:10s}  {len(text):>7,}자")
        except Exception as e:
            errors.append({'filename': f.name, 'error': str(e)})
            print(f"  ERR {f.name}: {e}", file=sys.stderr)

    OUT.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding='utf-8')
    print(f"\nExtracted {len(results)} docs → {OUT}")
    print(f"Total chars: {sum(r['char_count'] for r in results):,}")
    if errors:
        print(f"Errors: {len(errors)}")
        for e in errors:
            print('  ', e)


if __name__ == '__main__':
    main()
